import re
import io
from numpy import True_
import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

st.title("Automated Resume Entity & Skill Extractor")

# ---------- Session states ----------
if "messages" not in st.session_state:
    st.session_state.messages = []
if "name" not in st.session_state:
    st.session_state.name = None
if "DOB" not in st.session_state:
    st.session_state.DOB = None
if "ADDR" not in st.session_state:
    st.session_state.ADDR = None
if "PHONE" not in st.session_state:
    st.session_state.PHONE = None
if "EMAIL" not in st.session_state:
    st.session_state.EMAIL = None
if "CAREER_OBJECTIVE" not in st.session_state:
    st.session_state.CAREER_OBJECTIVE = None
if "DEGREE" not in st.session_state:
    st.session_state.DEGREE = None    
if "PROJECTS" not in st.session_state:
    st.session_state.PROJECTS = None
if "SOFT_SKILLS" not in st.session_state:
    st.session_state.SOFT_SKILLS = None

if "PROF_SKILLS" not in st.session_state:
    st.session_state.PROF_SKILLS = None
if "APPLIED_POSITION" not in st.session_state:
    st.session_state.APPLIED_POSITION = None

if "TEMPLATE" not in st.session_state:
    st.session_state.TEMPLATE = None

if "docx_bytes" not in st.session_state:
    st.session_state.docx_bytes = None
if "step" not in st.session_state:
    # steps: "intro", "name", "dob", "done"
    st.session_state.step = "intro"

# ---------- Intro (only first time) ----------
if not st.session_state.messages:
    st.session_state.messages.append({
        "role": "assistant",
        "content": (
            "မင်္ဂလာပါ၊ ကျွန်တော်ကတော့ အလိုအလျောက် CV နဲ့ ကျွမ်းကျင်မှုတွေ ထုတ်ဖော်ရေးသားပေးသွားမှာပဲဖြစ်ပါတယ်။\n\n"
            "CV ရေးသားဖို့ ဘယ်ကနေ ဘယ်လိုစရေးရမှန်း မသိတဲ့ user တွေအတွက် အကူအညီပေးဖို့ ရည်ရွယ်ပါတယ်။ \n CV ကို တည်ဆောက်ဖို့အတွက် ကျွန်တော်က မေးခွန်းအချို့ကို မေးမြန်းသွားမှာပဲ ဖြစ်ပါတယ်။ \n"
            "သင်ဖြေခဲ့တဲ့ ကိုယ်ရေးကိုယ်တာ အချက်အလက်တွေကို ကျွန်တော်က သိမ်းထားမှာ မဟုတ်ကြောင်း ဦးစွာ အသိပေးပါရစေ။"
        )
    })
    st.session_state.step = "name"  # after intro, go to name question

# ---------- Helper methods ----------
def insert_hr(paragraph):
    p = paragraph._p  # <w:p> element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(
        pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')      # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def extract_name(text: str) -> str:
    if not text:
        return ""
    text = text.strip()

    patterns = [
        r"my name is\s+(.+)",   # My name is ...
        r"i am\s+(.+)",         # I am ...
        r"i'm\s+(.+)",           # I'm ...
        r".*ကတော့\s+(.+)\s+.*(ပဲဖြစ်ပါတယ်။|ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
        r".*က\s+(.+)\s+.*(ပဲဖြစ်ပါတယ်။|ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$"
    ]

    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            name = m.group(1).strip(". ").strip()
            return name

    return text.strip(". ")  # fallback: full text

def extract_dob(text: str) -> str:
    if not text:
        return ""
    text = text.strip(". ").strip()

    patterns = [
        r".*\sis\s+(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})$",
        r"(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})$",
        r".*ကတော့\s+(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
        r".*က\s+(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
        r".*ကတော့\s+(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})\s+.*",
        r".*က\s+(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})\s+.*",
        r"(0[1-9]|[12][0-9]|3[0-1])\/(0[1-9]|1[0-2])\/([0-9]{4})",
        r"(0[1-9]|[12][0-9]|3[0-1])-(0[1-9]|1[0-2])-([0-9]{4})",
        r"(0[1-9]|[12][0-9]|3[0-1]),(0[1-9]|1[0-2]),([0-9]{4})",
        r"(0[1-9]|[12][0-9]|3[0-1])\s(0[1-9]|1[0-2])\s([0-9]{4})"

    ]

    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            day = m.group(1).strip()
            month = m.group(2).strip()
            year = m.group(3).strip()

            match month:
                case "01": mon = "January"
                case "02": mon = "February"
                case "03": mon = "March"
                case "04": mon = "April"
                case "05": mon = "May"
                case "06": mon = "June"
                case "07": mon = "July"
                case "08": mon = "August"
                case "09": mon = "September"
                case "10": mon = "October"
                case "11": mon = "November"
                case "12": mon = "December"
                case _: mon = "Unknown Month"

            return f"{day} {mon} {year}"

    return text  # fallback

def extract_addr(text: str) -> str:
    if not text:
        return ""
    text = text.strip(". ").strip()

    patterns = [
        r".*\sis\s+(.+)\.$",
        r".*ကတော့\s+(.+)\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
        r".*က\s+(.+)\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
        r".*ကတော့\s+(.+)\s+မှာ.*",
        r".*က\s+(.+)\s+မှာ.*"
    ]

    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            addr = m.group(1).strip()

            return addr

    return text  # fallback

def extract_phone(text: str) -> str:
    if not text:
        return ""
    text = text.strip(". ").strip()

    patterns = [
        r".*\s(is)\s+09([0-9]{9})\.$",
        r"^(09)([0-9]{9})\.?$",
        r".*(ကတော့|က)\s+09([0-9]{9})\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
    ]

    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            phone = m.group(2).strip()

            return f"+95 9{phone}"

    return text  # fallback

def extract_email(text: str) -> str:
    if not text:
        return ""
    text = text.strip(". ").strip()

    patterns = [
        r".*\sis\s+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\.$",
        r".*(ကတော့|က)\s+([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\s+.*(ဖြစ်ပါတယ်။|ပါ။|ဖြစ်တယ်။)$",
    ]

    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE)
        if m:
            email = m.group(2).strip()

            return email

    return text  # fallback

def process_career_objective(text: str) -> str:
    if not text:
        return ""

    original = text.strip()
    lower = original.lower()

    # 1) Detect role / identity (very simple keyword-based)
    role = None
    if "student" in lower:
        role = "student"
    if "developer" in lower or "programmer" in lower or "engineer" in lower:
        role = "software developer"
    if "data" in lower and "analyst" in lower:
        role = "data analyst"
    if "business" in lower and "analyst" in lower:
            role = "business analyst"
    # Default if nothing obvious
    if role is None:
        role = "aspiring professional"

    st.session_state.APPLIED_POSITION = role

    # 2) Detect goal phrase
    goal = None
    trigger_phrases = [
        "want to", "would like to", "aim to", "aiming to", "plan to",
        "my goal is", "my objective is", "seeking", "looking for",
        "hope to", "hoping to"
    ]
    for phrase in trigger_phrases:
        idx = lower.find(phrase)
        if idx != -1:
            # Take everything after the phrase as goal clause
            start = idx + len(phrase)
            goal = original[start:].strip(" .")
            break

    # 3) Detect skills / strengths (simple keywords)
    skills_keywords = [
        "android", "flutter", "php", "mysql", "javascript",
        "python", "nlp", "machine learning", "data science",
        "teamwork", "leadership", "communication", "problem-solving", "odoo", "erp", "django", "react", "vue", "angular",
        "web development", "mobile development", "software development", "data analysis", "project management"
    ]
    skills_found = [kw for kw in skills_keywords if kw in lower]
    st.session_state.PROF_SKILLS = skills_found

    # Build skill/value sentence
    value_sentence = ""
    if skills_found:
        # Make them readable (capitalize first letter)
        pretty_skills = ", ".join(s.title() for s in skills_found)
        
        value_sentence = (
            f" I bring skills in {pretty_skills} and am eager to contribute to real-world projects."
        )

    # 4) Compose final objective (2–3 sentences)
    if goal:
        objective = (
            f"As a {role}, I am seeking opportunities to {goal}. "
            f"My aim is to grow professionally while adding value to my organization.{value_sentence}"
        )
    else:
        # No explicit goal phrase; keep it generic
        objective = (
            f"As a {role}, I am seeking opportunities to develop my skills and gain practical experience. "
            f"I am motivated to learn, grow, and contribute to my future team.{value_sentence}"
        )

    return objective

def translate_career_objective_myanmar_to_english(text: str) -> str:
    if not text:
        return ""

    replacements = {
        "ရည်မှန်းချက်": "goal",
        "အလုပ်": "job",
        "အလုပ်လျှောက်": "apply for a job",
        "အလုပ်ရ": "get a job",
        "စီမံကိန်း": "project",
        "တိုးတက်": "improve",
        "ကျွမ်းကျင်မှု": "skills",
        "အတွေ့အကြုံ": "experience",
        "လေ့လာ": "learn",
        "အသုံးချ": "apply",
        "ထည့်ဝင်": "contribute",
        "ဆောင်ရွက်": "work",
        "တည်ဆောက်": "build",
        "ပံ့ပိုး": "support",
        "ရယူ": "gain",
        "ဖွံ့ဖြိုး": "develop",
        "အသိပညာ": "knowledge",
        "ပညာရပ်": "field",
        "အနာဂတ်": "future",
        "တာဝန်": "responsibility",
        "ပရော်ဖက်ရှင်နယ်": "professional",
        "ဖွံ့ဖြိုးရေး": "development",
        "အဖွဲ့": "team",
        "ကုမ္ပဏီ": "company",
        "အင်ဂျင်နီယာ": "engineer",
        "developer": "developer",
        "software developer": "software developer",
        "ကျောင်းသား": "student",
        "ဆော့ဖ်ဝဲ": "software",
        "အင်တာနက်": "internet",
        "အသင်းအဖွဲ့": "team",
        "အသင်းအဖွဲ့ဝင်": "team member",
        "အဖွဲ့အစည်း": "organization",
        "အသင်းအဖွဲ့နဲ့ အလုပ်လုပ်နိုင်": "work well in a team",
        "လျင်မြန်စွာ လေ့လာနိုင်သူ": "quick learner",
        "လေ့လာနိုင်သူ": "learner",
        "ကိုယ့်ကိုယ်ကို အလိုအလျောက် အားပေးနိုင်သူ": "self-motivated",
        "ဖိအား": "pressure",
        "ဖိအားနှင့် အလုပ်လုပ်နိုင်သူ": "handles pressure well",
        "ပြောဆိုဆက်ဆံရေး ရေးကောင်းမွန်သူ": "communication skills",
        "ပြောဆိုဆက်ဆံရေး": "communication"
    }

    translated = text
    for burmese, english in replacements.items():
        translated = translated.replace(burmese, english)

    return translated

def extract_degree(text: str) -> str:
    if not text:
        return []

    results = []
    parts = [p.strip() for p in text.split(".") if p.strip()]

    for part in parts:
        match = re.match(r"^.*(ကတော့|က)\s+(.+?),\s*(.+?),\s*(\d{4})$", part)
        if match:
            results.append({
                "degree": match.group(2).strip(),
                "university": match.group(3).strip(),
                "year": match.group(4).strip()
            })
        else:
            results.append({"raw": part})

    return results

def extract_projects(text: str) -> str:
    if not text:
        return []

    if re.search(r"(မရှိ|မရှိသေးပါ|မရှိပါ|no)", text, flags=re.IGNORECASE):
        return "No projects"
    
    results = []
    parts = [p.strip() for p in text.split(".") if p.strip()]

    for part in parts:
        match = re.match(r"^(.+?)$", part)
        match1 = re.search(r"(မရှိ|မရှိသေးပါ|မရှိပါ|no)", part)
        if match:
            results.append(match.group(1).strip())        
        else:
            results.append({"raw": part})

    return results

def process_soft_skills(text: str) -> str:
    if not text:
        return ""

    original = text.strip()
    lower = original.lower()

    skills_keywords = [
        "quick learner", "team", "self-motivated", "under pressure", "communication skills", "learner", "handles pressure well", "communication"]

    skills_found = [kw for kw in skills_keywords if kw in lower]
    if skills_found:
        soft_skills = [
            "Quick learner with a strong interest in exploring new technologies.",
            "Self-motivated and able to work effectively as part of a team.",
            "Capable of performing well under pressure and meeting deadlines.",
            "Strong communication skills."
        ]

    return soft_skills

def extract_tp(text: str) -> str:
    if not text:
        return ""
    text = text.strip(". ").strip()

    # Match 1–3 in English or Myanmar digits
    pattern = r"([1-3]|[၁-၃])"

    m = re.search(pattern, text)
    if m:
        return m.group(1)

    return text  # fallback

def set_cell_background(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge, attrs in kwargs.items():
        element = OxmlElement(f"w:{edge}")
        for attr, val in attrs.items():
            element.set(qn(f"w:{attr}"), val)
        tc_borders.append(element)

def set_cell_all_borders(cell):
    """Add solid black borders on all 4 sides of a cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")        # ~0.75 pt
        element.set(qn("w:color"), "000000")
        tc_borders.append(element)

def replace_text_node(node, replacement):
    """
    Replace the text of a Word XML <w:t> node.

    Converts \\n into actual Word line breaks.
    """

    parent = node.getparent()

    if parent is None:
        return

    replacement = (
        ""
        if replacement is None
        else str(replacement)
    )

    # Split replacement into lines
    lines = replacement.split("\n")

    # ---------------------------------------------------------
    # Simple case: no line break required
    # ---------------------------------------------------------

    if len(lines) == 1:

        node.text = replacement

        return

    # ---------------------------------------------------------
    # Multiple lines
    # ---------------------------------------------------------

    # The first line stays in the existing <w:t>
    node.text = lines[0]

    current_run = parent

    for line in lines[1:]:

        # Add line break
        br = OxmlElement("w:br")
        current_run.append(br)

        # Add new text node
        new_text = OxmlElement("w:t")
        new_text.text = line

        current_run.append(new_text)


def replace_placeholders_in_document(doc, replacements):
    """
    Replace placeholders in:

        - normal document text
        - tables
        - text boxes
        - headers
        - footers

    \\n in replacement values becomes an actual
    Word line break.
    """

    parts = []

    # Main document
    parts.append(doc.element)

    # Headers and footers
    for section in doc.sections:

        parts.append(
            section.header._element
        )

        parts.append(
            section.footer._element
        )

    # ---------------------------------------------------------
    # Search all Word text nodes
    # ---------------------------------------------------------

    for part in parts:

        text_nodes = part.xpath(
            ".//*[local-name()='t']"
        )

        for node in list(text_nodes):

            if node.text is None:
                continue

            for placeholder, replacement in replacements.items():

                if placeholder not in node.text:
                    continue

                replacement = (
                    ""
                    if replacement is None
                    else str(replacement)
                )

                # -------------------------------------------------
                # Placeholder is entirely inside this text node
                # -------------------------------------------------

                if node.text == placeholder:

                    replace_text_node(
                        node,
                        replacement
                    )

                    continue

                # -------------------------------------------------
                # Placeholder is part of this text node
                # -------------------------------------------------

                if placeholder in node.text:

                    # Keep text before placeholder
                    before, after = node.text.split(
                        placeholder,
                        1
                    )

                    node.text = before

                    parent = node.getparent()

                    # Find the run containing this text node
                    run = parent

                    lines = replacement.split("\n")

                    # Insert replacement
                    for index, line in enumerate(lines):

                        if index > 0:

                            br = OxmlElement("w:br")
                            run.append(br)

                        if line:

                            new_text = OxmlElement("w:t")
                            new_text.text = line

                            run.append(new_text)

                    # Add text after placeholder
                    if after:

                        new_text = OxmlElement("w:t")
                        new_text.text = after

                        run.append(new_text)

# ---------- Render existing history ----------
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])

# ---------- Name step ----------
if st.session_state.step == "name" and st.session_state.name is None:
    # Ask only once
    question = "ပထမဆုံး သင့်ရဲ့ နာမည်အပြည့်အစုံကို အင်္ဂလိပ်လို ရေးပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == question
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": question
        })
        with st.chat_message("assistant"):
            st.write(question)

    user_input = st.chat_input("သင့်ရဲ့ နာမည်ကို ဒီမှာရေးပေးပါ...")

    if user_input:
        # Store user message
        st.session_state.messages.append({
            "role": "user",
            "content": user_input
        })

        # Process name
        name = extract_name(user_input)
        st.session_state.name = name

        with st.chat_message("user"):
            st.write(user_input)

        # Move to DOB step
        st.session_state.step = "dob"
        st.rerun()

# ---------- DOB step ----------
elif st.session_state.step == "dob" and st.session_state.DOB is None:
    dobQuestion = "အခု သင့်ရဲ့ မွေးသက္ကရာဇ်ကို (dd/mm/yyyy) ပုံစံဖြင့် အင်္ဂလိပ် ဂဏန်းများ ထည့်သွင်းပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == dobQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": dobQuestion
        })
        with st.chat_message("assistant"):
            st.write(dobQuestion)

    dobInput = st.chat_input("သင့်ရဲ့ မွေးသက္ကရာဇ်ကို ဒီမှာ ရေးပေးပါ...")

    if dobInput:
        st.session_state.messages.append({
            "role": "user",
            "content": dobInput
        })

        dob = extract_dob(dobInput)
        st.session_state.DOB = dob

        with st.chat_message("user"):
            st.write(dobInput)

        # All basic info collected
        st.session_state.step = "addr"
        st.rerun()

# ---------- Addr step ----------
elif st.session_state.step == "addr" and st.session_state.ADDR is None:
    addrQuestion = "အခု သင့်ရဲ့ နေရပ်လိပ်စာကို ထည့်သွင်းပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == addrQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": addrQuestion
        })
        with st.chat_message("assistant"):
            st.write(addrQuestion)

    addrInput = st.chat_input("သင့်ရဲ့ နေရပ်လိပ်စာကို ဒီမှာ ရေးပေးပါ...")

    if addrInput:
        st.session_state.messages.append({
            "role": "user",
            "content": addrInput
        })

        addr = extract_addr(addrInput)
        st.session_state.ADDR = addr

        with st.chat_message("user"):
            st.write(addrInput)

        # All basic info collected
        st.session_state.step = "phone"
        st.rerun()

elif st.session_state.step == "phone" and st.session_state.PHONE is None:
    phoneQuestion = "အခု သင့်ရဲ့ ဖုန်းနံပါတ်ကို ထည့်သွင်းပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == phoneQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": phoneQuestion
        })
        with st.chat_message("assistant"):
            st.write(phoneQuestion)

    phoneInput = st.chat_input("သင့်ရဲ့ ဖုန်းနံပါတ်ကို ဒီမှာ ရေးပေးပါ...")

    if phoneInput:
        st.session_state.messages.append({
            "role": "user",
            "content": phoneInput
        })

        phone = extract_phone(phoneInput)
        st.session_state.PHONE = phone

        with st.chat_message("user"):
            st.write(phoneInput)

        # All basic info collected
        st.session_state.step = "email"
        st.rerun()

elif st.session_state.step == "email" and st.session_state.EMAIL is None:
    emailQuestion = "အခု သင့်ရဲ့ အီးမေးလ်လိပ်စာကို ထည့်သွင်းပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == emailQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": emailQuestion
        })
        with st.chat_message("assistant"):
            st.write(emailQuestion)

    emailInput = st.chat_input("သင့်ရဲ့ အီးမေးလ်လိပ်စာကို ဒီမှာ ရေးပေးပါ...")

    if emailInput:
        st.session_state.messages.append({
            "role": "user",
            "content": emailInput
        })

        email = extract_email(emailInput)
        st.session_state.EMAIL = email

        with st.chat_message("user"):
            st.write(emailInput)

        # All basic info collected
        st.session_state.step = "career_objective"
        st.rerun()

elif st.session_state.step == "career_objective" and st.session_state.CAREER_OBJECTIVE is None:
    careerQuestion = "အခု သင့်ရဲ့ အလုပ်ရည်ရွယ်ချက် ကို ထည့်သွင်းပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == careerQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": careerQuestion
        })
        with st.chat_message("assistant"):
            st.write(careerQuestion)

    careerInput = st.chat_input("သင့်ရဲ့ အလုပ်ရည်ရွယ်ချက်ကို ဒီမှာ ရေးပေးပါ...")

    if careerInput:
        st.session_state.messages.append({
            "role": "user",
            "content": careerInput
        })

        career = ""
        if(re.search(r"[\u1000-\u109F\uAA60-\uAA7F]", careerInput)):
            career_translated = translate_career_objective_myanmar_to_english(careerInput)
            career = process_career_objective(career_translated)
        else: 
            career = process_career_objective(careerInput)

        st.session_state.CAREER_OBJECTIVE = career

        with st.chat_message("user"):
            st.write(careerInput)

        # All basic info collected
        st.session_state.step = "degree"
        st.rerun()

elif st.session_state.step == "degree" and st.session_state.DEGREE is None:
    degreeQuestion = "အခု သင်ရရှိထားသော ဘွဲ့ သို့မဟုတ် Diploma များ ရှိပါက ဘွဲ့ နာမည်၊ ရရှိတဲ့ တက္ကသိုလ် (သို့) ကောလိပ်၊ ရရှိခဲ့သည့် ခုနှစ်ကို ပြောပြပေးပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == degreeQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": degreeQuestion
        })
        with st.chat_message("assistant"):
            st.write(degreeQuestion)

    degreeInput = st.chat_input("သင့်ရဲ့ ဘွဲ့များကို ဒီမှာ ရေးပေးပါ...")

    if degreeInput:
        st.session_state.messages.append({
            "role": "user",
            "content": degreeInput
        })

        degrees = extract_degree(degreeInput)

        st.session_state.DEGREE = degrees

        with st.chat_message("user"):
            st.write(degreeInput)

        # All basic info collected
        st.session_state.step = "projects"
        st.rerun()

elif st.session_state.step == "projects" and st.session_state.PROJECTS is None:
    prjQuestion = "အခု သင်ပြုလုပ်ထားသော အလုပ်များ၊ ပရောဂျက်များ ရှိပါက ဖော်ပြပါ။"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == prjQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": prjQuestion
        })
        with st.chat_message("assistant"):
            st.write(prjQuestion)

    prjInput = st.chat_input("သင့်ရဲ့ ပရောဂျက်များကို ဒီမှာ ရေးပေးပါ...")

    if prjInput:
        st.session_state.messages.append({
            "role": "user",
            "content": prjInput
        })

        projects = extract_projects(prjInput)

        st.session_state.PROJECTS = projects

        with st.chat_message("user"):
            st.write(prjInput)

        # All basic info collected
        st.session_state.step = "soft_skills"
        st.rerun()

elif st.session_state.step == "soft_skills" and st.session_state.SOFT_SKILLS is None:
    softSkillsQuestion = "အခု သင့်မှာရှိတဲ့ soft skills များကို ဖော်ပြပေးပါ။ (ဥပမာ - လျင်မြန်စွာ သင်ယူနိုင်သူ၊ အသင်းအဖွဲ့နှင့် လုပ်နိုင်သူ၊ ပြောဆိုဆက်ဆံရေး ကောင်းမွန်သူ)"

    # Ask only once: check for exact question text in history
    if not any(
        m["role"] == "assistant" and m["content"] == softSkillsQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": softSkillsQuestion
        })
        with st.chat_message("assistant"):
            st.write(softSkillsQuestion)

    softSkillsInput = st.chat_input("သင့်ရဲ့ soft skills များကို ဒီမှာ ရေးပေးပါ...")

    if softSkillsInput:
        st.session_state.messages.append({
            "role": "user",
            "content": softSkillsInput
        })

        soft_skills = ""
        if(re.search(r"[\u1000-\u109F\uAA60-\uAA7F]", softSkillsInput)):
            soft_translated = translate_career_objective_myanmar_to_english(softSkillsInput)
            soft_skills = process_soft_skills(soft_translated)
        else: 
            soft_skills = process_soft_skills(softSkillsInput)

        st.session_state.SOFT_SKILLS = soft_skills

        with st.chat_message("user"):
            st.write(softSkillsInput)

        # All basic info collected
        st.session_state.step = "template"
        st.rerun()

elif st.session_state.step == "template" and st.session_state.TEMPLATE is None:
    templateQuestion = (
    "မိမိလိုချင်တဲ့ cv ဖောင် ပုံစံကို ရွေးပေးပါ။ "
    "(ဥပမာ - ပုံစံ ၁, ပုံစံ ၂)"
    )

    # Show question only once
    if not any(
        m["role"] == "assistant" and m["content"] == templateQuestion
        for m in st.session_state.messages
    ):
        st.session_state.messages.append({
            "role": "assistant",
            "content": templateQuestion
        })

    # Render messages
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.write(msg["content"])

            # Add 3 expandable images for this question
            if msg["role"] == "assistant" and msg["content"] == templateQuestion:
                col1, col2, col3 = st.columns(3)

                with col1:
                    with st.expander("ပုံစံ ၁"):
                        st.image(
                            "img/template_1.png"                            
                        )

                with col2:
                    with st.expander("ပုံစံ ၂"):
                        st.image(
                            "img/template_2.png"
                        )

                with col3:
                    with st.expander("ပုံစံ ၃"):
                        st.image(
                            "img/template_3.png"
                        )

    # Input box
    tpInput = st.chat_input("သင့် project အမျိုးအစားကို ဒီမှာ ရေးပေးပါ...")

    if tpInput:
        st.session_state.messages.append({
            "role": "user",
            "content": tpInput
        })

        tp = extract_tp(tpInput)

        st.session_state.TEMPLATE = tp

        with st.chat_message("user"):
            st.write(tpInput)

        # All basic info collected
        st.session_state.step = "done"
        st.rerun()


# ---------- Generate Word file ----------
if st.session_state.step == "done" and st.session_state.docx_bytes is None:
    with st.chat_message("assistant"):
        st.write(
            "ကျွန်တော် သင့်ရဲ့ CV ကို ထုတ်ပေးနေပါသဖြင့် ခေတ္တခဏ စောင့်ပါ ခင်ဗျာ။"
        )

    progress_bar = st.progress(0)
    for pct in range(0, 101, 25):
        progress_bar.progress(pct)

    doc = Document()

    if st.session_state.TEMPLATE == "2" or st.session_state.TEMPLATE == "၂":
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titleRun = title.add_run("Curriculum Vitae")
        titleRun.bold = True
        titleRun.underline = True
        titleRun.font.size = Pt(36)
        titleRun.font.name = "Times New Roman"

        introtable = doc.add_table(rows=5, cols=2)
        introtable.autofit = False
        introtable.allow_autofit = False

        col_widths = [Cm(4), Cm(11)]  # adjust as needed        

        data = [
            ["Name", st.session_state.name],
            ["Date of Birth", st.session_state.DOB],
            ["Phone Number", st.session_state.PHONE],
            ["Email Address", st.session_state.ADDR],
            ["Contact Address", st.session_state.EMAIL]
        ]

        for i, row_data in enumerate(data):
            for j, value in enumerate(row_data):
                introtable.cell(i, j).text = value

        # Color first column of each row (including header if you want)
        for row in introtable.rows:
            first_cell = row.cells[0]
            set_cell_background(first_cell, "D9E1F2")
            for paragraph in first_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True  

        for row in introtable.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "4", "color": "000000"}
                )

        for i, width in enumerate(col_widths):
            for cell in introtable.columns[i].cells:
                cell.width = width

        contextTable = doc.add_table(rows=10, cols=1)
        contextTable.autofit = False
        contextTable.allow_autofit = False

        contextTable.columns[0].width = Cm(15)

        # Row 1: title
        title_cell = contextTable.rows[0].cells[0]
        title_run = title_cell.paragraphs[0].add_run("Educational Background")
        title_run.bold = True
        title_run.font.size = Pt(12)
        title_run.font.name = "Times New Roman"
        set_cell_background(title_cell, "D9E1F2")

        # Row 2: bullet list of values
        content_cell = contextTable.rows[1].cells[0]
        # Clear default paragraph if you want
        content_cell.paragraphs[0].clear()

        for res in st.session_state.DEGREE:
            p = content_cell.add_paragraph()
            p.style = "List Bullet"  # built-in bullet style

            text = f"{res['degree']}, {res['university']}, {res['year']}"
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Row 1: title
        exp_title_cell = contextTable.rows[2].cells[0]
        exp_title_run = exp_title_cell.paragraphs[0].add_run("Experience (Projects, Internships, Training Course, etc...)")
        exp_title_run.bold = True
        exp_title_run.font.size = Pt(12)
        exp_title_run.font.name = "Times New Roman"
        set_cell_background(exp_title_cell, "D9E1F2")

        # Row 2: bullet list of values
        exp_content_cell = contextTable.rows[3].cells[0]
        # Clear default paragraph if you want
        exp_content_cell.paragraphs[0].clear()

        for res in st.session_state.PROJECTS:
            p = exp_content_cell.add_paragraph()
            p.style = "List Bullet"  # built-in bullet style

            text = res
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Row 1: title
        p_skills_title_cell = contextTable.rows[4].cells[0]
        p_skills_title_run = p_skills_title_cell.paragraphs[0].add_run("Technical Skills")
        p_skills_title_run.bold = True
        p_skills_title_run.font.size = Pt(12)
        p_skills_title_run.font.name = "Times New Roman"
        set_cell_background(p_skills_title_cell, "D9E1F2")

        # Row 2: bullet list of values
        p_skills_content_cell = contextTable.rows[5].cells[0]
        # Clear default paragraph if you want
        p_skills_content_cell.paragraphs[0].clear()

        for res in st.session_state.PROF_SKILLS:
            p = p_skills_content_cell.add_paragraph()
            p.style = "List Bullet"  # built-in bullet style

            text = res
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Row 1: title
        s_skills_title_cell = contextTable.rows[6].cells[0]
        s_skills_title_run = s_skills_title_cell.paragraphs[0].add_run("Soft Skills(Communication, Collaboration, Education, etc...)")
        s_skills_title_run.bold = True
        s_skills_title_run.font.size = Pt(12)
        s_skills_title_run.font.name = "Times New Roman"
        set_cell_background(s_skills_title_cell, "D9E1F2")

        # Row 2: bullet list of values
        s_skills_content_cell = contextTable.rows[7].cells[0]
        # Clear default paragraph if you want
        s_skills_content_cell.paragraphs[0].clear()

        for res in st.session_state.SOFT_SKILLS:
            p = s_skills_content_cell.add_paragraph()
            p.style = "List Bullet"  # built-in bullet style

            text = res
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.name = "Times New Roman"

        # Row 1: title
        pos_title_cell = contextTable.rows[8].cells[0]
        pos_title_run = pos_title_cell.paragraphs[0].add_run("Applied Position")
        pos_title_run.bold = True
        pos_title_run.font.size = Pt(12)
        pos_title_run.font.name = "Times New Roman"
        set_cell_background(pos_title_cell, "D9E1F2")
        # Row 2: bullet list of values
        pos_content_cell = contextTable.rows[9].cells[0]
        # Clear default paragraph if you want
        pos_content_cell.paragraphs[0].clear()
        
        p = pos_content_cell.add_paragraph()
        text = st.session_state.APPLIED_POSITION
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Times New Roman"

        for row in contextTable.rows:
            for cell in row.cells:
                set_cell_all_borders(cell)


    elif st.session_state.TEMPLATE == "3" or st.session_state.TEMPLATE == "၃":

        doc = Document("temp/tp1.docx")

        # Education
        education_lines = []

        for res in st.session_state.DEGREE:

            education_lines.append(
                f"{res['degree']}    \t {res['university']}    \t {res['year']}"
            )

        education_text = "\n".join(
            education_lines
        )

        # Projects
        if st.session_state.PROJECTS == "No projects":
            projects_text = ""
        else:
            projects_text = "\n".join(
                str(project)
                for project in st.session_state.PROJECTS
            )

        # Technical Skills
        technical_skills_text = "\n".join(
            str(skill)
            for skill in st.session_state.PROF_SKILLS
        )

        # Soft Skills
        soft_skills_text = "\n".join(
            str(skill)
            for skill in st.session_state.SOFT_SKILLS
        )

        # Placeholder mapping
        replacements = {

            "{{APP_POS}}":
                st.session_state.APPLIED_POSITION,

            "{{NAME}}":
                st.session_state.name,

            "{{DOB}}":
                st.session_state.DOB,

            "{{PHONE}}":
                st.session_state.PHONE,

            "{{ADDR}}":
                st.session_state.ADDR,

            "{{EMAIL}}":
                st.session_state.EMAIL,

            "{{CAR_OBJ}}":
                st.session_state.CAREER_OBJECTIVE,

            "{{EDU}}":
                education_text,

            "{{PROJECTS}}":
                projects_text,

            "{{TECH_SKILLS}}":
                technical_skills_text,

            "{{SOFT_SKILLS}}":
                soft_skills_text,
        }

        # Replace placeholders, including text boxes
        replace_placeholders_in_document(
            doc,
            replacements
        )

                
    else:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titleRun = title.add_run("Curriculum Vitae")
        titleRun.bold = True
        titleRun.underline = True
        titleRun.font.size = Pt(36)
        titleRun.font.name = "Times New Roman"

        nameP = doc.add_paragraph()
        nameRun = nameP.add_run(st.session_state.name)
        nameRun.bold = True
        nameRun.font.size = Pt(24)
        nameRun.font.name = "Times New Roman"

        dobP = doc.add_paragraph()
        dobBullet = dobP.add_run("\u2666")
        dobBullet.bold = True
        dobBullet.font.size = Pt(12)
        dobBullet.font.name = "Segoe UI Symbol"
        dobLable = dobP.add_run("Date of Birth\t: ")
        dobLable.bold = True
        dobLable.font.size = Pt(12)
        dobLable.font.name = "Times New Roman"
        dobValue = dobP.add_run(st.session_state.DOB)    
        dobValue.font.size = Pt(12)
        dobValue.font.name = "Times New Roman"

        addrP = doc.add_paragraph()
        addrBullet = addrP.add_run("\u2666")
        addrBullet.bold = True
        addrBullet.font.size = Pt(12)
        addrBullet.font.name = "Segoe UI Symbol"
        addrLable = addrP.add_run("Address\t\t: ")
        addrLable.bold = True
        addrLable.font.size = Pt(12)
        addrLable.font.name = "Times New Roman"
        addrValue = addrP.add_run(st.session_state.ADDR)    
        addrValue.font.size = Pt(12)
        addrValue.font.name = "Times New Roman"

        phoneP = doc.add_paragraph()
        phoneBullet = phoneP.add_run("\u2666")
        phoneBullet.bold = True
        phoneBullet.font.size = Pt(12)
        phoneBullet.font.name = "Segoe UI Symbol"
        phoneLable = phoneP.add_run("Phone\t\t: ")
        phoneLable.bold = True
        phoneLable.font.size = Pt(12)
        phoneLable.font.name = "Times New Roman"
        phoneValue = phoneP.add_run(st.session_state.PHONE)
        phoneValue.font.size = Pt(12)
        phoneValue.font.name = "Times New Roman"

        emailP = doc.add_paragraph()
        emailBullet = emailP.add_run("\u2666")
        emailBullet.bold = True
        emailBullet.font.size = Pt(12)
        emailBullet.font.name = "Segoe UI Symbol"
        emailLable = emailP.add_run("Email\t\t: ")
        emailLable.bold = True
        emailLable.font.size = Pt(12)
        emailLable.font.name = "Times New Roman"
        emailValue = emailP.add_run(st.session_state.EMAIL)
        emailValue.font.size = Pt(12)
        emailValue.font.name = "Times New Roman"

        hr1P = doc.add_paragraph()
        insert_hr(hr1P)

        cbP = doc.add_paragraph()
        cbRun = cbP.add_run("Career Objective")
        cbRun.bold = True
        cbRun.underline = True
        cbRun.font.size = Pt(16)
        cbRun.font.name = "Times New Roman"

        cbValueP = doc.add_paragraph()
        cbValueRun = cbValueP.add_run(st.session_state.CAREER_OBJECTIVE)
        cbValueRun.font.size = Pt(12)
        cbValueRun.font.name = "Times New Roman"
        
        hr2P = doc.add_paragraph()
        insert_hr(hr2P)

        degP = doc.add_paragraph()
        degRun = degP.add_run("Academic Qualifications")
        degRun.bold = True
        degRun.underline = True
        degRun.font.size = Pt(16)
        degRun.font.name = "Times New Roman"

        for res in st.session_state.DEGREE:
            degreeP = doc.add_paragraph()        
            degreeValue = degreeP.add_run(res["degree"] + '\t' + res["university"] + '\t' + res["year"])
            degreeValue.font.size = Pt(11)
            degreeValue.font.name = "Times New Roman"

        hr2P = doc.add_paragraph()
        insert_hr(hr2P)

        if st.session_state.PROJECTS != "No projects":
            prjP = doc.add_paragraph()
            prjRun = prjP.add_run("Academic Projects")
            prjRun.bold = True
            prjRun.underline = True
            prjRun.font.size = Pt(16)
            prjRun.font.name = "Times New Roman"

            for res in st.session_state.PROJECTS:
                prjP = doc.add_paragraph()
                prjValue = prjP.add_run(res)
                prjValue.font.size = Pt(11)
                prjValue.font.name = "Times New Roman"

            hr2P = doc.add_paragraph()
            insert_hr(hr2P)

        softP = doc.add_paragraph()
        softRun = softP.add_run("Personal Skills")
        softRun.bold = True
        softRun.underline = True
        softRun.font.size = Pt(16)
        softRun.font.name = "Times New Roman"

        for res in st.session_state.SOFT_SKILLS:
            softV = doc.add_paragraph()
            softBullet = softV.add_run("\u2666\t")
            softBullet.bold = True
            softBullet.font.size = Pt(12)
            softBullet.font.name = "Segoe UI Symbol"
            softValue = softV.add_run(res)
            softValue.font.size = Pt(11)
            softValue.font.name = "Times New Roman"

        hr2P = doc.add_paragraph()
        insert_hr(hr2P)

        psP = doc.add_paragraph()
        psRun = psP.add_run("Professional Skills")
        psRun.bold = True
        psRun.underline = True
        psRun.font.size = Pt(16)
        psRun.font.name = "Times New Roman"

        for res in st.session_state.PROF_SKILLS:
            psV = doc.add_paragraph()
            psBullet = psV.add_run("\u2666\t")
            psBullet.bold = True
            psBullet.font.size = Pt(12)
            psBullet.font.name = "Segoe UI Symbol"
            psValue = psV.add_run(res)
            psValue.font.size = Pt(11)
            psValue.font.name = "Times New Roman"

        hr2P = doc.add_paragraph()
        insert_hr(hr2P)

        applP = doc.add_paragraph()
        applRun = applP.add_run("Applied Position")
        applRun.bold = True
        applRun.underline = True
        applRun.font.size = Pt(16)
        applRun.font.name = "Times New Roman"

        applValueP = doc.add_paragraph()
        applValueRun = applValueP.add_run(st.session_state.APPLIED_POSITION)
        applValueRun.font.size = Pt(12)
        applValueRun.font.name = "Times New Roman"

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.session_state.docx_bytes = buffer.read()

    with st.chat_message("assistant"):
        st.write(
            "သင့်ရဲ့ CV ကို ဒေါင်းလုဒ် လုပ်နိုင်ပါပြီ။ အောက်က ခလုတ်ကို နှိပ်ကာ ဒေါင်းလုဒ် လုပ်နိုင်ပါတယ်။ သင်ဖြေသမျှ ကိုယ်ရေးအချက်အလက်များကို ကျွန်တော် သိမ်းထားမည် မဟုတ်ကြောင်း ထပ်မံ အသိပေးပါရစေ"
        )

# ---------- Download button ----------
if st.session_state.docx_bytes is not None:
    st.download_button(
        label="Download CV",
        data=st.session_state.docx_bytes,
        file_name="cv.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )